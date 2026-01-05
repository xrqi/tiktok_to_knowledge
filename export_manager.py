from typing import List, Dict, Any, Optional
from pathlib import Path
import json
import markdown
from datetime import datetime
import os

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


class ExportManager:
    """Export Manager - Handles exporting knowledge in various formats"""
    
    def __init__(self):
        self.supported_formats = ['json', 'markdown', 'txt']
        if DOCX_AVAILABLE:
            self.supported_formats.append('docx')
        if REPORTLAB_AVAILABLE:
            self.supported_formats.append('pdf')
    
    def export_knowledge_to_format(self, knowledge_points: List[Dict[str, Any]], 
                                 output_path: str, format_type: str = 'json') -> bool:
        """Export knowledge points to specified format"""
        try:
            output_path = Path(output_path)
            
            # Ensure output directory exists
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            if format_type.lower() == 'json':
                return self._export_to_json(knowledge_points, output_path)
            elif format_type.lower() == 'markdown' or format_type.lower() == 'md':
                return self._export_to_markdown(knowledge_points, output_path)
            elif format_type.lower() == 'txt':
                return self._export_to_txt(knowledge_points, output_path)
            elif format_type.lower() == 'docx':
                if not DOCX_AVAILABLE:
                    print("python-docx not available. Please install with: pip install python-docx")
                    return False
                return self._export_to_docx(knowledge_points, output_path)
            elif format_type.lower() == 'pdf':
                if not REPORTLAB_AVAILABLE:
                    print("reportlab not available. Please install with: pip install reportlab")
                    return False
                return self._export_to_pdf(knowledge_points, output_path)
            else:
                print(f"Unsupported format: {format_type}. Supported formats: {self.supported_formats}")
                return False
                
        except Exception as e:
            print(f"Error exporting to {format_type}: {e}")
            return False
    
    def _export_to_json(self, knowledge_points: List[Dict[str, Any]], output_path: Path) -> bool:
        """Export knowledge points to JSON format"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(knowledge_points, f, ensure_ascii=False, indent=2)
            return True
        except Exception as e:
            print(f"Error exporting to JSON: {e}")
            return False
    
    def _export_to_markdown(self, knowledge_points: List[Dict[str, Any]], output_path: Path) -> bool:
        """Export knowledge points to Markdown format"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                # Add title and metadata
                f.write(f"# Knowledge Export\n\n")
                f.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                
                for i, kp in enumerate(knowledge_points, 1):
                    f.write(f"## {i}. {kp.get('title', 'Untitled')}\n\n")
                    f.write(f"**Category:** {kp.get('category', 'Uncategorized')}\n\n")
                    f.write(f"**Importance:** {kp.get('importance', 3)}/5\n\n")
                    
                    if 'tags' in kp and kp['tags']:
                        f.write(f"**Tags:** {', '.join(kp['tags'])}\n\n")
                    
                    f.write(f"**Content:**\n{kp.get('content', '')}\n\n")
                    
                    if i < len(knowledge_points):
                        f.write("---\n\n")
            return True
        except Exception as e:
            print(f"Error exporting to Markdown: {e}")
            return False
    
    def _export_to_txt(self, knowledge_points: List[Dict[str, Any]], output_path: Path) -> bool:
        """Export knowledge points to plain text format"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write("Knowledge Export\n")
                f.write("=" * 50 + "\n")
                f.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                
                for i, kp in enumerate(knowledge_points, 1):
                    f.write(f"{i}. {kp.get('title', 'Untitled')}\n")
                    f.write("-" * 50 + "\n")
                    f.write(f"Category: {kp.get('category', 'Uncategorized')}\n")
                    f.write(f"Importance: {kp.get('importance', 3)}/5\n")
                    
                    if 'tags' in kp and kp['tags']:
                        f.write(f"Tags: {', '.join(kp['tags'])}\n")
                    
                    f.write(f"\nContent:\n{kp.get('content', '')}\n\n")
                    f.write("=" * 50 + "\n\n")
            return True
        except Exception as e:
            print(f"Error exporting to TXT: {e}")
            return False
    
    def _export_to_docx(self, knowledge_points: List[Dict[str, Any]], output_path: Path) -> bool:
        """Export knowledge points to Word (docx) format"""
        try:
            doc = Document()
            
            # Add title
            title_para = doc.add_paragraph()
            title_run = title_para.add_run("Knowledge Export")
            title_run.bold = True
            title_run.font.size = docx.shared.Pt(16)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add generation date
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph()  # Empty paragraph for spacing
            
            for i, kp in enumerate(knowledge_points, 1):
                # Add knowledge point title
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(f"{i}. {kp.get('title', 'Untitled')}")
                title_run.bold = True
                title_run.font.size = docx.shared.Pt(12)
                
                # Add metadata
                doc.add_paragraph(f"Category: {kp.get('category', 'Uncategorized')}")
                doc.add_paragraph(f"Importance: {kp.get('importance', 3)}/5")
                
                if 'tags' in kp and kp['tags']:
                    doc.add_paragraph(f"Tags: {', '.join(kp['tags'])}")
                
                # Add content
                doc.add_paragraph("Content:")
                doc.add_paragraph(kp.get('content', ''))
                
                # Add separator if not the last item
                if i < len(knowledge_points):
                    doc.add_paragraph()  # Empty paragraph for spacing
                    doc.add_paragraph("-" * 50)
                    doc.add_paragraph()  # Empty paragraph for spacing
            
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error exporting to DOCX: {e}")
            return False
    
    def _export_to_pdf(self, knowledge_points: List[Dict[str, Any]], output_path: Path) -> bool:
        """Export knowledge points to PDF format"""
        try:
            doc = SimpleDocTemplate(str(output_path), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Add title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            title = Paragraph("Knowledge Export", title_style)
            story.append(title)
            
            # Add generation date
            date_para = Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal'])
            story.append(date_para)
            story.append(Spacer(1, 12))
            
            for i, kp in enumerate(knowledge_points, 1):
                # Add knowledge point title
                title_para = Paragraph(f"{i}. {kp.get('title', 'Untitled')}", styles['Heading2'])
                story.append(title_para)
                
                # Add metadata
                story.append(Paragraph(f"<b>Category:</b> {kp.get('category', 'Uncategorized')}", styles['Normal']))
                story.append(Paragraph(f"<b>Importance:</b> {kp.get('importance', 3)}/5", styles['Normal']))
                
                if 'tags' in kp and kp['tags']:
                    story.append(Paragraph(f"<b>Tags:</b> {', '.join(kp['tags'])}", styles['Normal']))
                
                # Add content
                story.append(Paragraph("<b>Content:</b>", styles['Heading3']))
                content_para = Paragraph(kp.get('content', ''), styles['Normal'])
                story.append(content_para)
                
                # Add separator if not the last item
                if i < len(knowledge_points):
                    story.append(Spacer(1, 20))
                    story.append(Paragraph("-" * 50, styles['Normal']))
                    story.append(Spacer(1, 20))
            
            doc.build(story)
            return True
        except Exception as e:
            print(f"Error exporting to PDF: {e}")
            return False

    def batch_export(self, knowledge_points: List[Dict[str, Any]], 
                    base_path: str, formats: List[str]) -> Dict[str, bool]:
        """Export knowledge points to multiple formats"""
        results = {}
        base_path = Path(base_path)
        
        for fmt in formats:
            output_path = base_path.with_suffix(f'.{fmt}')
            success = self.export_knowledge_to_format(knowledge_points, str(output_path), fmt)
            results[fmt] = success
        
        return results


# Example usage
if __name__ == "__main__":
    # Example knowledge points
    knowledge_points = [
        {
            "title": "Artificial Intelligence",
            "content": "AI is a branch of computer science that aims to create software or machines that exhibit human-like intelligence. This includes learning from experience, understanding natural language, solving problems, and recognizing patterns.",
            "category": "Technology",
            "tags": ["AI", "Machine Learning", "Computer Science"],
            "importance": 5
        },
        {
            "title": "Machine Learning",
            "content": "Machine learning is a method of data analysis that automates analytical model building. It's a branch of artificial intelligence based on the idea that systems can learn from data, identify patterns and make decisions with minimal human intervention.",
            "category": "Technology", 
            "tags": ["ML", "Algorithms", "Data Science"],
            "importance": 4
        }
    ]
    
    exporter = ExportManager()
    print(f"Supported formats: {exporter.supported_formats}")
    
    # Export to different formats
    success_json = exporter.export_knowledge_to_format(knowledge_points, "output/knowledge_export.json", "json")
    success_md = exporter.export_knowledge_to_format(knowledge_points, "output/knowledge_export.md", "markdown")
    success_txt = exporter.export_knowledge_to_format(knowledge_points, "output/knowledge_export.txt", "txt")
    
    print(f"JSON export: {'Success' if success_json else 'Failed'}")
    print(f"Markdown export: {'Success' if success_md else 'Failed'}")
    print(f"TXT export: {'Success' if success_txt else 'Failed'}")